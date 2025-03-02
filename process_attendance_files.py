import os
import sys

script_dir = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, script_dir)

from tools import write_attendance_sheet, is_chinese,register_chinese_font, get_pdf_styles

import pandas as pd
from openpyxl import Workbook
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# word转pdf在除了windows，其他平台都有点麻烦且效果不好，干脆多搞点代码直接生成pdf                                 
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.units import mm, inch


def process_attendance_files(data, date, year, month, day, output_folder):

    data["姓名是否中文"] = data["姓名"].apply(is_chinese)

    # 筛选数据
    filtered_data_cn = data[(data["旷课课时"] >= 2) & data["姓名是否中文"]]
    filtered_data_cn = filtered_data_cn.sort_values(by=["旷课课时", "旷课次数"], ascending=[False, False])

    filtered_data_int = data[(data["旷课课时"] >= 2) & ~data["姓名是否中文"]]
    filtered_data_int = filtered_data_int.sort_values(by=["旷课课时", "旷课次数"], ascending=[False, False])

    # 考勤表excel
    def create_excel():
        wb = Workbook()
        
        # 本科生 sheet
        stu_cn = wb.active
        stu_cn.title = "本科生"
        cols = ["学号", "姓名", "学院", "班级", "旷课次数", "迟到次数", "早退次数", "旷课课时"]
        write_attendance_sheet(stu_cn, filtered_data_cn, cols, [15, 10, 10, 10, 10, 10, 10, 10], ["学院", "班级"])
        
        # 留学生 sheet
        international = wb.create_sheet(title="留学生")
        write_attendance_sheet(international, filtered_data_int, cols, [15, 35, 10, 10, 10, 10, 10, 10], ["学院", "班级"])

        # 保存文件
        excel_output = f"{output_folder}/计算机科学与技术学院学生第{date}上课啦系统缺勤情况.xlsx"
        wb.save(excel_output)

    # 考勤通报word文档
    def create_docx():
        # 创建 Word 文档
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'  # 必须先设置font.name
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        title = f"第{date}“上课啦”考勤通报"
        para_1 = doc.add_paragraph()
        run_1 = para_1.add_run(title)
        run_1.font.size = Pt(22)
        run_1.font.bold = True
        para_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 设置标题居中

        # 段落内容
        content = f"以下同学在第{date}“上课啦”考勤中未正常出勤, 旷课学时计入个人档案, 并纳入日常考评。"
        para_2 = doc.add_paragraph()
        run_2 = para_2.add_run(content)
        run_2.font.size = Pt(16)
        para_2.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 设置内容两端对齐

        # 添加表格
        table = doc.add_table(rows=1, cols=6, style="Table Grid")
        col_width_dict = {0: 1.6, 1: 1.12, 2: 0.7638, 3: 0.7638, 4: 0.7638, 5: 0.7638}
        row_height = Pt(25)
        
        # 设置列宽
        for col_num in range(6):
            table.cell(0, col_num).width = Inches(col_width_dict[col_num])

        # 设置表头
        headers = ["学号", "姓名", "旷课次数", "迟到次数", "早退次数", "旷课课时"]
        header_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            cell = header_cells[idx]
            cell.text = header                      # 设置单元格文本
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True            # 设置文字加粗

        # 添加数据行
        for _, row in filtered_data_cn.iterrows():
            row_cells = table.add_row().cells
            row_cells[0].text = str(row["学号"])
            row_cells[1].text = row["姓名"]
            row_cells[2].text = str(row["旷课次数"])
            row_cells[3].text = str(row["迟到次数"])
            row_cells[4].text = str(row["早退次数"])
            row_cells[5].text = str(row["旷课课时"])

        # 全局设置行高和垂直对齐方式
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = 1         # 单元格垂直居中
                row.height = row_height             # 设置行高
                for paragraph in cell.paragraphs:
                    paragraph.alignment = 1         # 设置单元格文本居中
                    for run in paragraph.runs:
                        run.font.size = Pt(11)

        # 添加说明
        note_1 = "\n注:\n一、根据学生手册中“浙江师范大学学生违纪处分规定”中第三章第二十七条规定,学生一学期内旷课累计满10学时的, 给予警告处分; 满20学时的, 给予严重警告处分;满30学时的, 给与记过处分; 满40学时的, 给与留校察看处分; 因旷课屡次受到纪律处分并经教育不改的, 可以给予开除学籍处分;\n二、学时计算方法如下: \n(1) 旷课1小节为1学时, 未经请假缺勤1天, 不足5学时按5学时计; 超过5学时的, 按实际学时数计; \n(2) 学生无故迟到或早退达3次, 作旷课1学时计; "

        para_3 = doc.add_paragraph()
        run_3 = para_3.add_run(note_1)
        run_3.font.size = Pt(12)
        run_3.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # 设置内容两端对齐

        note_3 = "特此通报！"
        para_4 = doc.add_paragraph()
        run_4 = para_4.add_run(note_3)
        run_4.font.size = Pt(16)
        run_4.font.bold = True
        run_4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # 设置内容左对齐

        para_5 = doc.add_paragraph()
        para_5.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # 设置段落右对齐
        run_5 = para_5.add_run(f"计算机科学与技术学院学工办\n{year}年{month}月{day}日")
        run_5.font.size = Pt(16)

        # 保存文件
        docx_output = f"{output_folder}/计算机科学与技术学院学生第{date}上课啦系统缺勤通报.docx"
        doc.save(docx_output)
        return docx_output

    # 考勤通报PDF（直接生成）
    def create_pdf():
        # 注册中文字体
        register_chinese_font()
        cn_font = "SimSun"
        styles_dict = get_pdf_styles()

        # 创建PDF文档
        pdf_output = f"{output_folder}/计算机科学与技术学院学生第{date}上课啦系统缺勤通报.pdf"
        doc = SimpleDocTemplate(
            pdf_output,
            pagesize=A4,
            leftMargin=20*mm,
            rightMargin=20*mm,
            topMargin=15*mm,
            bottomMargin=15*mm
        )

        # 构建内容元素
        elements = []

        # 标题
        title = f"第{date}“上课啦”考勤通报"
        elements.append(Paragraph(title, styles_dict["title"]))

        # 正文内容
        content = f"以下同学在第{date}“上课啦”考勤中未正常出勤, 旷课学时计入个人档案, 并纳入日常考评。"
        elements.append(Paragraph(content, styles_dict["content"]))

        # 创建表格
        headers = ["学号", "姓名", "旷课次数", "迟到次数", "早退次数", "旷课课时"]
        table_data = [[Paragraph(h, styles_dict["table_header"]) for h in headers]]
        
        for _, row in filtered_data_cn.iterrows():
            table_row = [
                Paragraph(str(row["学号"]), styles_dict['table_cell']),
                Paragraph(row["姓名"], styles_dict['table_cell']),
                Paragraph(str(row["旷课次数"]), styles_dict['table_cell']),
                Paragraph(str(row["迟到次数"]), styles_dict['table_cell']),
                Paragraph(str(row["早退次数"]), styles_dict['table_cell']),
                Paragraph(str(row["旷课课时"]), styles_dict['table_cell'])
            ]
            table_data.append(table_row)

        # 设置列宽
        effective_width = (A4[0] - 46*mm) / inch  # 有效宽度（英寸）
        original_widths = [1.3, 1.22, 0.8038, 0.8038, 0.8038, 0.8038]  # 列宽比例，数字有点抽象不过问题不大
        scale = effective_width / sum(original_widths)
        col_widths = [w * scale * inch for w in original_widths]
        
        # 创建表格对象
        table = Table(table_data, colWidths=col_widths, rowHeights=10*mm)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.lightgrey),
            ('BOLD', (0,0), (-1,0), True),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
            ('FONTNAME', (0,0), (-1,0), cn_font),
            ('FONTSIZE', (0,0), (-1,-1), 11),
            ('GRID', (0,0), (-1,-1), 0.5, colors.black),
            ('BOX', (0,0), (-1,-1), 0.5, colors.black),
        ]))
        elements.append(table)
        elements.append(Spacer(1, 10*mm))

        # 添加说明
        note_text = """<para>
        <font size=12>注:</font><br/>
        <font size=12>一、根据学生手册中“浙江师范大学学生违纪处分规定”中第三章第二十七条规定，学生一学期内旷课累计满10学时的，给予警告处分；满20学时的，给予严重警告处分；满30学时的，给与记过处分；满40学时的，给与留校察看处分；因旷课屡次受到纪律处分并经教育不改的，可以给予开除学籍处分；</font><br/>
        <font size=12>二、学时计算方法如下：</font><br/>
        <font size=12>(1) 旷课1小节为1学时，未经请假缺勤1天，不足5学时按5学时计；超过5学时的，按实际学时数计；</font><br/>
        <font size=12>(2) 学生无故迟到或早退达3次，作旷课1学时计；</font>
        </para>"""
        elements.append(Paragraph(note_text, styles_dict['note']))

        # 特此通报
        elements.append(Paragraph("特此通报！", styles_dict['sign']))

        # 日期和学院
        date_text = f"计算机科学与技术学院学工办<br/>{year}年{month}月{day}日"
        elements.append(Paragraph(date_text, styles_dict['date']))

        # 生成PDF
        doc.build(elements)
        return pdf_output

    # 执行
    create_excel()
    create_docx()
    create_pdf()